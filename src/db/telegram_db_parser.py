import sqlite3
import hashlib
from docx import Document
from parser import clean_message  # ✅ 텔레그램 메시지 정리 함수
from keywords_list import find_keywords  # ✅ 키워드 리스트 로드

# ✅ 데이터베이스 및 DOCX 파일 설정
DB_FILE = "db/telegram_messages.db"
DOCX_FILE = "db/telegram_messages.docx"

def create_database(db_file):
    """📌 데이터베이스 및 테이블 생성"""
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sender TEXT,
            message TEXT,
            detected_keywords TEXT,  -- 감지된 키워드 저장
            hash TEXT UNIQUE
        )
    """)
    
    conn.commit()
    conn.close()
    print(f"✅ 데이터베이스 {db_file} 생성 완료!")

def get_hash(sender, message):
    """📌 데이터 중복 확인을 위한 해시(MD5) 생성"""
    combined_text = (sender + message).encode('utf-8')
    return hashlib.md5(combined_text).hexdigest()

def insert_messages(messages, db_file):
    """
    📌 텔레그램 메시지를 데이터베이스에 저장하고,
    삽입된 데이터를 기준으로 키워드를 검출함.
    """
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()

    inserted_count = 0
    skipped_count = 0
    all_messages = []
    keyword_detected_list = []

    for msg in messages:
        sender = msg.get("sender", "알 수 없음")
        raw_message = msg.get("message", "")
        cleaned_message = clean_message(raw_message)

        hash_value = get_hash(sender, cleaned_message)

        # ✅ 중복 메시지 확인
        cursor.execute("SELECT id FROM messages WHERE hash = ?", (hash_value,))
        if cursor.fetchone():
            skipped_count += 1
            continue  # 중복 메시지는 저장하지 않음

        # ✅ 메시지 저장
        cursor.execute("INSERT INTO messages (sender, message, detected_keywords, hash) VALUES (?, ?, ?, ?)",
                       (sender, cleaned_message, "N/A", hash_value))
        conn.commit()
        inserted_count += 1

        # ✅ 키워드 감지
        detected_keywords = find_keywords(cleaned_message)
        keyword_str = ", ".join(detected_keywords) if detected_keywords else "N/A"

        # ✅ 키워드가 감지되었다면 데이터 업데이트
        cursor.execute("UPDATE messages SET detected_keywords = ? WHERE hash = ?", (keyword_str, hash_value))
        conn.commit()

        # ✅ 모든 메시지를 리스트에 추가
        all_messages.append({
            "sender": sender,
            "message": cleaned_message,
            "keywords": detected_keywords if detected_keywords else ["N/A"]
        })

        # ✅ 키워드 감지된 메시지 저장
        if detected_keywords:
            keyword_detected_list.append({
                "sender": sender,
                "message": cleaned_message,
                "keywords": detected_keywords
            })

    conn.close()
    
    print(f"✅ {db_file}: {inserted_count}개의 메시지가 저장되었습니다. (중복 {skipped_count}건 제외)")
    
    return all_messages, keyword_detected_list

def export_to_docx(messages, docx_file):
    """📌 키워드가 감지된 메시지만 DOCX 파일로 저장"""
    keyword_detected_data = [msg for msg in messages if msg.get("keywords") and msg["keywords"] != ["N/A"]]

    if not keyword_detected_data:
        print(f"❌ {docx_file}에 저장할 키워드 감지 메시지가 없습니다.")
        return

    doc = Document()
    doc.add_heading("📄 Telegram Messages Report (Keywords Detected)", level=1)

    for msg in keyword_detected_data:
        doc.add_heading(msg["sender"], level=2)
        doc.add_paragraph(f"📝 Message: {msg['message']}")
        doc.add_paragraph(f"📌 Keywords Detected: {', '.join(msg['keywords'])}")
        doc.add_paragraph("=" * 50)

    doc.save(docx_file)
    print(f"📄 {docx_file} (키워드 감지 메시지만) 파일 저장 완료!")

# ✅ 데이터베이스 생성 실행
if __name__ == "__main__":
    create_database(DB_FILE)  # ✅ DB_FILE을 인자로 전달하여 생성
